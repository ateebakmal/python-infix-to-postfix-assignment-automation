from docx import Document

    # The following code will add dark borders to each cell in the table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
class PostixAssignmentSolver():
    
    @staticmethod
    def __set_cell_border(cell, **kwargs):
        """
        Set borders for a table cell.

        Args:
        cell: The cell to apply the border.
        kwargs: Accepts 'top', 'left', 'bottom', 'right', where each is a dict containing
                'sz' (size), 'val' (border style), 'color' (hex color code), 'space' (space size).
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Create a new set of borders
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            if border_name in kwargs:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), kwargs[border_name].get('val', 'single'))
                border.set(qn('w:sz'), str(kwargs[border_name].get('sz', 12)))  # size of border
                border.set(qn('w:space'), str(kwargs[border_name].get('space', 0)))
                border.set(qn('w:color'), kwargs[border_name].get('color', '000000'))  # black color
                tcBorders.append(border)

        tcPr.append(tcBorders)

    @staticmethod
    def __make_complete_table_dark(table):
        for row in table.rows:
            for cell in row.cells:
                PostixAssignmentSolver.__set_cell_border(cell,
                                top={'sz': 24, 'val': 'single', 'color': '000000'},
                                bottom={'sz': 24, 'val': 'single', 'color': '000000'},
                                left={'sz': 24, 'val': 'single', 'color': '000000'},
                                right={'sz': 24, 'val': 'single', 'color': '000000'})
                
    
    @staticmethod
    def __precedence(op):
        if op in ('+', '-'):
            return 1
        if op in ('*', '/'):
            return 2
        if op == "^":
            return 3
        return 0

    @staticmethod
    #this function will return a list of all the outputs and takes parameters [current_row_number, current_character, postfix_list, stack]
    def __generateRowData(row_num, current_char, postfix_list, stack):
        stack_output = ""
        if(not stack):
            stack_output = "empty"
        else:
            stack_output = "".join(stack)
        return [str(row_num), (current_char), "".join(postfix_list) , stack_output]

    
    @staticmethod
    def __writeRowInTable(row_num, current_char, postfix_list, stack , table):
        row_data = PostixAssignmentSolver.__generateRowData(row_num, current_char , postfix_list, stack)
        for col in range(0,4):
            table.cell(row_num,col).text = row_data[col]

        return row_num + 1
    
    @staticmethod
    def __solveAndCreateWordDoc(index, infix_expression):

        #Make a doc
        doc = Document()
        #Make a table
        table = doc.add_table(1,4)

        # adding the headings
        headings = ["Step No", "Expression", "Postfix", "Stack"]
        for i in range(0,4):
            cell = table.cell(0, i)  # Access the cell in the first row (index 0)
            cell.text = headings[i]  # Set the text for each cell
            # Make the text bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        #----------------------------------------------------Main logic for solving the infix expression
        stack = []
        postfix = []
        row_num = 1
        for i in infix_expression:
            #Make a new row 
            table.add_row()

            if(i.isalnum()):
                postfix.append(i)

                row_num = PostixAssignmentSolver.__writeRowInTable(row_num, i, postfix, stack, table)
            elif(i == "("):
                stack.append(i)
                row_num = PostixAssignmentSolver.__writeRowInTable(row_num, i, postfix, stack, table)
            elif(i == ")"):
                while(stack and stack[len(stack) - 1] != "("):
                    postfix.append(stack[-1])
                    stack.pop()

                #pop ( from stack as well
                if(stack):
                    stack.pop()
                
                row_num = PostixAssignmentSolver.__writeRowInTable(row_num, i, postfix, stack, table)

            elif(stack and PostixAssignmentSolver.__precedence(i) > PostixAssignmentSolver.__precedence(stack[-1])): #if precedence is greater than top
                stack.append(i)

                row_num = PostixAssignmentSolver.__writeRowInTable(row_num, i, postfix, stack, table)

            else:
                while(stack and PostixAssignmentSolver.__precedence(stack[-1]) >= PostixAssignmentSolver.__precedence(i)):
                    postfix.append(stack[-1])
                    stack.pop()
                stack.append(i)

                row_num = PostixAssignmentSolver.__writeRowInTable(row_num, i, postfix, stack, table)

        while stack:
            postfix.append(stack[-1])
            stack.pop()

        table.add_row()
        row_num = PostixAssignmentSolver.__writeRowInTable(row_num, i, postfix, stack, table)

        #Making all the borders dark    
        PostixAssignmentSolver.__make_complete_table_dark(table)
        #saving the doc
        doc.save(f"postfix-question-'{index}'.docx")


    @staticmethod
    #Following function will solve all the prefix questions:
    def solveInfixToPostfix(infix_list):

        for index, question in enumerate(infix_list):
            print(f"Solving question {(index+1)}")
            PostixAssignmentSolver.__solveAndCreateWordDoc(index+1,question)
            print(f"Question{index+1} solved")

    

